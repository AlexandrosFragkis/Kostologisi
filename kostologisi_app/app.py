try:
    import streamlit as st
except ModuleNotFoundError:
    raise ImportError("Το Streamlit δεν είναι εγκατεστημένο. Εκτέλεσε 'pip install streamlit' για να συνεχίσεις.")

import json
import os
import pandas as pd
from io import BytesIO
from docx import Document
from PIL import Image

st.set_page_config(layout="wide", page_title="Κοστολόγηση Επίπλων", page_icon="📐")

# --- Custom CSS με αισθητική από masterwood.gr ---
st.markdown("""
    <style>
        .stApp {
            background-color: #f8f6f3;
            background-image: linear-gradient(to bottom, #1e2022, #3b1c1f);
            color: #fdfdfd;
            font-family: 'Segoe UI', sans-serif;
        }

        h1, h2, h3, h4, h5, h6 {
            color: #f5f1ee;
        }

        .stSidebar {
            background-color: #2d2d2d;
        }

        .stNumberInput input, .stTextInput input {
            background-color: #fff;
            color: #000;
            border-radius: 6px;
            padding: 0.25rem;
        }

        .stSelectbox>div>div>div {
            background-color: #fff;
            color: #000;
            border-radius: 6px;
        }

        .stButton>button {
            background-color: #8a1c1c;
            color: white;
            font-weight: bold;
            border-radius: 10px;
            padding: 0.5rem 1.5rem;
            transition: 0.3s;
        }

        .stButton>button:hover {
            background-color: #aa2e2e;
            transform: scale(1.03);
        }

        .block-container {
            padding-top: 2rem;
            padding-bottom: 2rem;
        }

        .css-1d391kg {
            padding: 2rem;
            border-radius: 12px;
            background-color: rgba(255, 255, 255, 0.05);
            box-shadow: 0 4px 12px rgba(0,0,0,0.2);
        }
    </style>
""", unsafe_allow_html=True)

# --- Αρχεία για αποθήκευση τιμών ---
PRICE_FILE = "material_prices.json"

# --- Προεπιλεγμένες Τιμές Υλικών για Κοστολόγηση ---
def_material_prices = {
    "Καπλαμάς Δρυς": 130,
    "Λάκα": 95,
    "Μελαμίνη": 60,
    "Duropal": 85,
    "Compact": 100,
    "Corian": 300,
    "Εξαρτήματα Επίπλων": 150
}

# --- Τιμές Αναφοράς Επίπλων ---
def_furniture_reference = {
    "Ντουλάπα ανοιγόμενη": {
        "Μελαμίνη": 320,
        "Λάκα ματ/σατινέ": 400,
        "Λάκα ζαγρέ": 430,
        "Λάκα ματ/σατινέ με ταμπλά": 550,
        "Λάκα ματ/σατινέ με πηχάκια": 430,
        "Καπλαμάς δρυς": 430,
        "Καπλαμάς δρυς με ταμπλά": 580,
        "Καπλαμάς δρυς με πηχάκια": 460,
        "Καπλαμάς καρυδιά": 520,
        "Καπλαμάς καρυδιά με ταμπλα": 670,
        "Καπλαμάς καρυδιά με πηχάκια": 550,
        "Τζάμι με μεταλλικό πλαίσιο": 550
    },
    "Ντουλάπα συρόμενη": {
        "Μελαμίνη": 400,
        "Λάκα ματ/σατινέ": 480,
        "Λάκα ζαγρέ": 520,
        "Καπλαμάς δρυς": 550,
        "Καπλαμάς καρυδιά": 600,
        "Με καθρέπτη": 480,
        "Τζάμι με μεταλλικό πλαίσιο": 580
    },
    "Ντουλάπα ανοιγόμενη με ταπετσαρία": {
        "Λάκα ματ/σατινέ": 600,
        "Καπλαμάς δρυς": 600,
        "Καπλαμάς καρυδιά": 700
    },
    "Κουζίνα": {
        "Πάγκος/Πλάτη Duropal 60 εκ.": 110,
        "Πάγκος Duropal 90 εκ.": 150,
        "Μελαμίνη πάνω": 350,
        "Μελαμίνη κάτω": 400,
        "Λάκα ματ/σατινέ πάνω": 400,
        "Λάκα ματ/σατινέ κάτω": 450,
        "Λάκα ζαγρέ πάνω": 450,
        "Λάκα ζαγρέ κάτω": 500,
        "Καπλαμάς δρυς πάνω": 450,
        "Καπλαμάς δρυς κάτω": 500
    },
    "Επενδύσεις (απλή)": {
        "Μελαμίνη": 60,
        "Λάκα ματ/σατινέ": 90,
        "Λάκα ζαγρέ": 100,
        "Καπλαμάς δρυς": 100,
        "Καπλαμάς καρυδιά": 150
    },
    "Επενδύσεις (πηχάκια)": {
        "Μελαμίνη": 150,
        "Λάκα ματ/σατινέ": 200,
        "Λάκα ζαγρέ": 250,
        "Καπλαμάς δρυς": 250,
        "Καπλαμάς καρυδιά": 300
    },
    "Επιπλέον Υλικά αναφοράς": {
        "Duropal": 110,
        "Compact": 150,
        "Corian": 600,
        "Εξαρτήματα επίπλων": 300
    }
}

# --- Αποθήκευση / Φόρτωση ---
def load_prices():
    if os.path.exists(PRICE_FILE):
        with open(PRICE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return def_material_prices.copy()

def save_prices(prices):
    with open(PRICE_FILE, "w", encoding="utf-8") as f:
        json.dump(prices, f, ensure_ascii=False, indent=2)

if "material_prices" not in st.session_state:
    st.session_state.material_prices = load_prices()
if "furniture_reference_prices" not in st.session_state:
    st.session_state.furniture_reference_prices = def_furniture_reference.copy()
if "furniture_list" not in st.session_state:
    st.session_state.furniture_list = []

# --- Sidebar ---
st.sidebar.header("🔧 Τιμές Υλικών")
with st.sidebar:
    st.subheader("Τιμές Υλικών Κοστολόγησης (€ / m² ή τεμ)")
    updated_prices = {}
    for mat in def_material_prices:
        val = st.number_input(f"{mat}", value=float(st.session_state.material_prices.get(mat, def_material_prices[mat])), min_value=0.0, key=f"price_{mat}")
        updated_prices[mat] = val
    if updated_prices != st.session_state.material_prices:
        st.session_state.material_prices = updated_prices
        save_prices(updated_prices)

    st.subheader("Τιμές Αναφοράς Επίπλων")
    for section, materials in st.session_state.furniture_reference_prices.items():
        with st.expander(section):
            for mat, price in materials.items():
                st.session_state.furniture_reference_prices[section][mat] = st.number_input(
                    f"{mat}",
                    value=float(price),
                    min_value=0.0,
                    key=f"ref_{section}_{mat}"
                )

# --- Κύριο Περιεχόμενο ---
st.title("📐 Κοστολόγηση Custom Επίπλων")

st.header("1. Εισαγωγή Διαστάσεων Κατασκευής (σε εκατοστά)")
construction_name = st.text_input("Όνομα Κατασκευής")
exterior_length = st.number_input("Μήκος σε cm", min_value=0.0, step=1.0)
exterior_height = st.number_input("Ύψος σε cm", min_value=0.0, step=1.0)
exterior_area = round((exterior_length * exterior_height) / 10000, 2)

st.header("2. Επιλογή Υλικών")
reference_options = []
for section, materials in st.session_state.furniture_reference_prices.items():
    for m in materials:
        reference_options.append(f"{section} - {m}")
selected_material = st.selectbox("Επιλογή Τύπου Κατασκευής & Υλικού", reference_options)
panel_material = st.selectbox("Υλικό Πάγκου (κουζίνα/μπάνιο)", options=["Τίποτα", "Compact", "Corian", "Duropal"])
add_hardware = st.selectbox("Προσθήκη Εξαρτημάτων Επίπλου", options=["ΟΧΙ", "ΝΑΙ"])

st.header("3. Συρτάρια")
drawer_count = st.number_input("Διαφορά από 4 συρτάρια (θετικός ή αρνητικός αριθμός)", value=0, step=1)
drawer_price = 250

total_drawers_cost = drawer_count * drawer_price

if st.button("Προσθήκη Επίπλου"):
    section, mat = selected_material.split(" - ", 1)
    price_per_m2 = st.session_state.furniture_reference_prices.get(section, {}).get(mat, 0)
    total_material_cost = exterior_area * price_per_m2
    panel_cost = 0 if panel_material == "Τίποτα" else st.session_state.material_prices.get(panel_material, 0)
    hardware_cost = st.session_state.material_prices["Εξαρτήματα Επίπλων"] if add_hardware == "ΝΑΙ" else 0
    total_cost = total_material_cost + panel_cost + hardware_cost + total_drawers_cost

    st.session_state.furniture_list.append({
        "Κατασκευή": construction_name,
        "Επιφάνεια (m²)": exterior_area,
        "Κατηγορία/Υλικό": selected_material,
        "Πάγκος": panel_material,
        "Εξαρτήματα": add_hardware,
        "Διαφορά Συρταριών": drawer_count,
        "Κόστος Υλικού": total_material_cost,
        "Κόστος Πάγκου": panel_cost,
        "Κόστος Εξαρτημάτων": hardware_cost,
        "Κόστος Συρταριών": total_drawers_cost,
        "Σύνολο": total_cost
    })
    st.success("✅ Προστέθηκε έπιπλο στη λίστα")

if st.session_state.furniture_list:
    st.subheader("📋 Λίστα Επίπλων")
    df = pd.DataFrame(st.session_state.furniture_list)
    st.dataframe(df, use_container_width=True)

    total = df["Σύνολο"].sum()
    st.success(f"💰 Συνολικό Κόστος Όλων των Επίπλων: {total:.2f} €")

    if st.button("Επαναφορά Λίστας"):
        st.session_state.furniture_list = []
        st.rerun()

    doc = Document()
    doc.add_heading("Λίστα Επίπλων", level=1)
    for item in st.session_state.furniture_list:
        for k, v in item.items():
            doc.add_paragraph(f"{k}: {v}")
        doc.add_paragraph("---")

    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)

    st.download_button(
        label="📥 Λήψη Λίστας σε Word",
        data=word_buffer,
        file_name="lista_epiplon.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.header("4. Επιπλέον Υπολογισμοί")
manual_cost = st.number_input("Χειροκίνητο συνολικό κόστος κατασκευής (€)", min_value=0.0, step=10.0)
commission_percent = st.number_input("Ποσοστό προμήθειας αρχιτέκτονα (%)", min_value=0.0, max_value=100.0, step=1.0)
commission_amount = manual_cost * (commission_percent / 100)
final_cost = manual_cost + commission_amount

st.write(f"🔧 Προμήθεια ({commission_percent:.0f}%): {commission_amount:.2f} €")
if manual_cost > 0:
    st.success(f"🏁 Τελικό Κόστος με Προμήθεια: {final_cost:.2f} €")
